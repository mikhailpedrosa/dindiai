from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from typing import Dict, Any
import os

class Exporter:
    def export_markdown(self, report_text: str, filename: str = "relatorio_financeiro.md"):
        """
        Exporta o texto para Markdown.
        """
        with open(filename, "w", encoding="utf-8") as f:
            f.write(report_text)
        return filename

    def export_docx(self, report_text: str, filename: str = "relatorio_financeiro.docx"):
        """
        Exporta o texto para DOCX.
        """
        doc = Document()
        doc.add_heading('Relatório Agente de Finanças Dindiai', 0)
        
        # Simples quebra de linhas e parágrafos
        for line in report_text.split('\n'):
            if line.startswith('##'):
                doc.add_heading(line.replace('##', '').strip(), level=1)
            elif line.startswith('###'):
                doc.add_heading(line.replace('###', '').strip(), level=2)
            else:
                doc.add_paragraph(line)
        
        doc.save(filename)
        return filename

    def export_pdf(self, report_text: str, filename: str = "relatorio_financeiro.pdf"):
        """
        Exporta o texto para PDF (Simples).
        """
        c = canvas.Canvas(filename, pagesize=letter)
        width, height = letter
        
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, height - 50, "Relatório Agente de Finanças Dindiai")
        
        c.setFont("Helvetica", 10)
        y = height - 80
        for line in report_text.split('\n'):
            if y < 50:
                c.showPage()
                y = height - 50
                c.setFont("Helvetica", 10)
            
            # Limitar largura da linha (simplificação)
            clean_line = line.strip()[:100]
            c.drawString(50, y, clean_line)
            y -= 15
            
        c.save()
        return filename
